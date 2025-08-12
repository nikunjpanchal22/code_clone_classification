def combine_word_documents(files) :
	combined_document = Document('empty.docx')
	count, number_of_files = 0, len(files)
	for file in files :
		sub_doc = Document(file)
		if count < number_of_files - 1 :
			sub_doc.add_page_break()
		for element in sub_doc._document_part.body._element :
			combined_document._document_part.body._element.append(element)
		count += 1
	combined_document.save('combined_word_documents.docx')





from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
def combine_word_documents(files):
    newdoc = Document()
    for i in range(len(files)):
        doc = Document(files[i])
        newdoc.add_paragraph("Document" + str(i+1), style='Title')
        for para in doc.paragraphs:
            text = para.text
            newdoc.add_paragraph(text, style='BodyText')
        newdoc.add_page_break()


